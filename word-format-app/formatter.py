from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

TITLE_FONT = "方正小标宋_GBK"
BODY_FONT = "方正仿宋_GBK"
LEVEL1_FONT = "方正黑体_GBK"
LEVEL2_FONT = "方正楷体_GBK"
WESTERN_FONT = "Times New Roman"
PAGE_FONT = "宋体"

TITLE_SIZE = Pt(22)
BODY_SIZE = Pt(16)
PAGE_SIZE = Pt(14)
LINE_SPACING = Pt(29.5)

DOC_NO_RE = re.compile(r".+〔\d{4}〕\d+号$")
LEVEL1_RE = re.compile(r"^[一二三四五六七八九十]+、")
LEVEL2_RE = re.compile(r"^（[一二三四五六七八九十]+）")
LEVEL3_RE = re.compile(r"^\d+\.")
LEVEL4_RE = re.compile(r"^（\d+）")
ARABIC_ITEM_RE = re.compile(r"^\s*(\d{1,2})([.．、]?)(\s*.+)?$")
CN_BRACKET_ITEM_RE = re.compile(r"^\s*[（(]([一二三四五六七八九十]+)[）)](\s*.+)?$")
DATE_RE = re.compile(r"^\d{4}\s*年\s*\d{1,2}\s*月(?:\s*\d{1,2}\s*日)?$")
SIGNATURE_DATE_RE = re.compile(
    r"^(?:\d{4}|[_＿]{2,})\s*年\s*(?:\d{1,2}|[_＿]{1,})\s*月(?:\s*(?:\d{1,2}|[_＿]{1,})\s*日)?$"
)
ATTACHMENT_RE = re.compile(r"^附件[：:]")
LATIN_RE = re.compile(r"[A-Za-z0-9]")
KEYPOINT_RE = re.compile(r"(一是|二是|三是|四是|五是|六是|七是|八是|九是|十是)")
CN_NUMERALS = "一二三四五六七八九十"
ZH_CHAR_RE = r"\u4e00-\u9fff"
YEAR_PREFIX_RE = re.compile(r"^\s*(?:19|20)\d{2}\s*年(?:度)?")


@dataclass
class ParagraphReport:
    index: int
    kind: str
    text: str
    font_east_asia: str
    font_ascii: str
    font_size_pt: float
    alignment: str
    first_line_indent_pt: float
    left_indent_pt: float
    line_spacing_pt: float
    actions: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass
class FormatSummary:
    paragraphs_updated: int = 0
    tables_updated: int = 0
    footer_updated: bool = False
    warnings: list[str] = field(default_factory=list)
    paragraph_reports: list[ParagraphReport] = field(default_factory=list)

    def to_dict(self) -> dict:
        payload = asdict(self)
        payload["warning_count"] = len(self.warnings)
        payload["paragraph_count"] = len(self.paragraph_reports)
        payload["kind_counts"] = _count_kinds(self.paragraph_reports)
        return payload


def format_document(input_path: Path, output_path: Path) -> FormatSummary:
    doc = Document(str(input_path))
    summary = FormatSummary()

    _set_document_layout(doc)
    _normalize_text_spacing(doc)
    _normalize_paragraph_structure(doc)
    _split_heading_body_paragraphs(doc)
    _normalize_inline_subheadings(doc)
    _normalize_numbered_sequences(doc)
    _format_paragraphs(doc, summary)
    _format_tables(doc, summary)
    summary.footer_updated = _set_page_footer(doc)

    if summary.footer_updated:
        summary.warnings.append("页码已设置为双面打印外侧页码：奇数页在右侧，偶数页在左侧。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return summary


def _count_kinds(paragraph_reports: list[ParagraphReport]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for item in paragraph_reports:
        counts[item.kind] = counts.get(item.kind, 0) + 1
    return counts


def _set_document_layout(doc: DocumentObject) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(32)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(15)
        section.start_type = WD_SECTION.CONTINUOUS


def _normalize_paragraph_structure(doc: DocumentObject) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text
        if not text or "\n" not in text:
            continue

        parts = [part.strip() for part in text.splitlines() if part.strip()]
        if len(parts) <= 1:
            continue

        paragraph.text = parts[0]
        anchor = paragraph
        for part in parts[1:]:
            anchor = _insert_paragraph_after(anchor, part)


def _normalize_inline_subheadings(doc: DocumentObject) -> None:
    paragraphs = doc.paragraphs
    idx = 0
    while idx < len(paragraphs) - 1:
        current = paragraphs[idx]
        following = paragraphs[idx + 1]
        current_text = current.text.strip()
        following_text = following.text.strip()

        if (
            _is_inline_subheading_candidate(current_text)
            and _is_body_following_paragraph(following_text)
        ):
            current.text = f"{current_text}{following_text}"
            _remove_paragraph(following)
            paragraphs = doc.paragraphs
            continue
        idx += 1


def _split_heading_body_paragraphs(doc: DocumentObject) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        split_parts = _split_level2_heading_and_body(text)
        if split_parts is None:
            continue

        heading, body = split_parts
        paragraph.text = heading
        _insert_paragraph_after(paragraph, body)


def _split_level2_heading_and_body(text: str) -> tuple[str, str] | None:
    if not LEVEL2_RE.match(text):
        return None

    match = re.match(r"^(（[一二三四五六七八九十]+）.+?[。！？：:])(.+)$", text)
    if not match:
        return None

    heading = match.group(1).strip()
    body = match.group(2).strip()
    if not body or not _is_body_following_paragraph(body):
        return None

    if len(heading) > 32:
        return None

    return heading, body


def _is_inline_subheading_candidate(text: str) -> bool:
    if not text or len(text) > 40:
        return False
    return bool((LEVEL3_RE.match(text) or LEVEL4_RE.match(text)) and text.endswith(("。", "？", "！")))


def _is_body_following_paragraph(text: str) -> bool:
    return bool(
        text
        and not LEVEL1_RE.match(text)
        and not LEVEL2_RE.match(text)
        and not LEVEL3_RE.match(text)
        and not LEVEL4_RE.match(text)
        and not ATTACHMENT_RE.match(text)
        and not DOC_NO_RE.match(text)
        and not DATE_RE.match(text)
        and not text.endswith(("：", ":"))
        and text.endswith(("。", "！", "？", "；"))
    )


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def _normalize_text_spacing(doc: DocumentObject) -> None:
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        normalized = _clean_spacing(paragraph.text)
        if normalized != paragraph.text:
            paragraph.text = normalized


def _clean_spacing(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        line = re.sub(r"[ \t]{2,}", " ", line)
        line = re.sub(rf"([{ZH_CHAR_RE}])\s+([{ZH_CHAR_RE}])", r"\1\2", line)
        line = re.sub(rf"([{ZH_CHAR_RE}])\s+([，。；：！？、】【（）《》“”‘’])", r"\1\2", line)
        line = re.sub(rf"([，。；：！？、】【（）《》“”‘’])\s+([{ZH_CHAR_RE}])", r"\1\2", line)
        line = re.sub(r"^([一二三四五六七八九十]+、)\s+", r"\1", line)
        line = re.sub(r"^(（[一二三四五六七八九十]+）)\s+", r"\1", line)
        line = re.sub(r"^(\d+\.)\s+", r"\1", line)
        line = re.sub(r"^(附件[：:])\s+", r"\1", line)
        line = re.sub(r"\s+([：:；，。！？])", r"\1", line)
        lines.append(line)
    return "\n".join(lines)


def _normalize_numbered_sequences(doc: DocumentObject) -> None:
    paragraphs = doc.paragraphs
    arabic_counter = None
    chinese_counter = None
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        if LEVEL1_RE.match(text):
            arabic_counter = None
            chinese_counter = None
            continue

        if _is_hard_sequence_boundary(text):
            arabic_counter = None
            chinese_counter = None
            continue

        if _is_chinese_bracket_candidate(text):
            if chinese_counter is not None or _has_following_chinese_bracket_item(paragraphs, idx):
                chinese_counter = 1 if chinese_counter is None else chinese_counter + 1
                paragraph.text = _replace_leading_cn_bracket_number(text, chinese_counter)
                _remove_word_numbering(paragraph)
            else:
                chinese_counter = None
            arabic_counter = None
            continue

        if _has_word_numbering(paragraph) or _is_arabic_item_candidate(text):
            if arabic_counter is not None or _has_following_arabic_item(paragraphs, idx):
                arabic_counter = 1 if arabic_counter is None else arabic_counter + 1
                paragraph.text = _replace_leading_number(text, arabic_counter)
                _remove_word_numbering(paragraph)
            else:
                if _has_word_numbering(paragraph):
                    _remove_word_numbering(paragraph)
                arabic_counter = None
            chinese_counter = None
            continue

        arabic_counter = None
        chinese_counter = None


def _is_hard_sequence_boundary(text: str) -> bool:
    return bool(
        DOC_NO_RE.match(text)
        or DATE_RE.match(text)
        or text.endswith(("：", ":"))
    )


def _has_word_numbering(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.numPr is not None)


def _remove_word_numbering(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.numPr
    if num_pr is not None:
        p_pr.remove(num_pr)


def _replace_leading_number(text: str, number: int) -> str:
    match = ARABIC_ITEM_RE.match(text)
    stripped = match.group(3).lstrip() if match and match.group(3) else text.lstrip()
    return f"{number}.{stripped}"


def _is_chinese_bracket_candidate(text: str) -> bool:
    if LEVEL1_RE.match(text):
        return False
    match = CN_BRACKET_ITEM_RE.match(text)
    return bool(match and match.group(2) and match.group(2).strip())


def _replace_leading_cn_bracket_number(text: str, number: int) -> str:
    match = CN_BRACKET_ITEM_RE.match(text)
    stripped = match.group(2).lstrip() if match and match.group(2) else text.lstrip()
    numeral = CN_NUMERALS[number - 1] if 1 <= number <= len(CN_NUMERALS) else str(number)
    return f"（{numeral}）{stripped}"


def _has_following_arabic_item(paragraphs: list[Paragraph], start_idx: int) -> bool:
    for paragraph in paragraphs[start_idx + 1:]:
        text = paragraph.text.strip()
        if not text:
            continue
        return _has_word_numbering(paragraph) or _is_arabic_item_candidate(text)
    return False


def _has_following_chinese_bracket_item(paragraphs: list[Paragraph], start_idx: int) -> bool:
    for paragraph in paragraphs[start_idx + 1:]:
        text = paragraph.text.strip()
        if not text:
            continue
        return _is_chinese_bracket_candidate(text)
    return False


def _is_arabic_item_candidate(text: str) -> bool:
    match = ARABIC_ITEM_RE.match(text)
    if not match:
        return False

    raw_number, punctuation, remainder = match.groups()
    if remainder is None:
        return False

    remainder = remainder.lstrip()
    if not remainder:
        return False

    if YEAR_PREFIX_RE.match(text) or DATE_RE.match(text):
        return False

    if punctuation:
        return True

    number = int(raw_number)
    if number > 10:
        return False

    if remainder.startswith(("年", "月", "日", "季度", "届", "次", "号")):
        return False

    return bool(re.match(rf"^[{ZH_CHAR_RE}“”\"'《（(]", remainder))


def _is_arabic_sequence_boundary(text: str) -> bool:
    return bool(
        _is_chinese_bracket_candidate(text)
        or LEVEL1_RE.match(text)
        or LEVEL4_RE.match(text)
        or ATTACHMENT_RE.match(text)
        or text.endswith(("：", ":"))
    )


def _format_paragraphs(doc: DocumentObject, summary: FormatSummary) -> None:
    title_indices = _detect_title_indices(doc.paragraphs)
    author_indices = _detect_author_indices(doc.paragraphs, title_indices)
    signature_indices, signature_date_index = _detect_signature_indices(doc.paragraphs)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            _normalize_paragraph_spacing(paragraph)
            continue

        kind = _classify_paragraph(
            text,
            idx,
            title_indices,
            author_indices,
            signature_indices,
            signature_date_index,
        )
        actions: list[str] = []
        warnings: list[str] = []

        if kind == "title":
            _normalize_title_text(paragraph)
            _apply_paragraph_style(
                paragraph,
                TITLE_FONT,
                TITLE_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True,
                keep_together=True,
            )
            actions.append("按标题设置为方正小标宋_GBK二号居中，自动优化断行")
        elif kind == "doc_number":
            _apply_paragraph_style(
                paragraph,
                BODY_FONT,
                BODY_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before_pt=LINE_SPACING.pt * 2,
            )
            actions.append("按发文字号设置为方正仿宋_GBK三号居中，下空2行")
            if "第" in text or "001" in text:
                warnings.append("发文字号可能包含“第”字或虚位序号，建议人工复核。")
        elif kind == "author":
            _apply_paragraph_style(
                paragraph,
                LEVEL2_FONT,
                BODY_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            actions.append("按作者名设置为方正楷体_GBK三号居中")
        elif kind == "signature":
            _apply_paragraph_style(
                paragraph,
                BODY_FONT,
                BODY_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )
            actions.append("按落款单位设置为方正仿宋_GBK三号右对齐")
        elif kind == "level1":
            _apply_paragraph_style(
                paragraph,
                LEVEL1_FONT,
                BODY_SIZE,
                bold=False,
                first_line_chars=2,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                keep_with_next=True,
                keep_together=True,
            )
            actions.append("按一级标题设置为方正黑体_GBK三号，首行缩进2字")
        elif kind == "level2":
            _apply_paragraph_style(
                paragraph,
                LEVEL2_FONT,
                BODY_SIZE,
                bold=False,
                first_line_chars=2,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                keep_with_next=True,
                keep_together=True,
            )
            actions.append("按二级标题设置为方正楷体_GBK三号，首行缩进2字")
        elif kind in {"level3", "level4", "attachment"}:
            _apply_paragraph_style(
                paragraph,
                BODY_FONT,
                BODY_SIZE,
                bold=False,
                first_line_chars=2,
                left_chars=2 if kind == "attachment" else None,
                space_before_pt=LINE_SPACING.pt if kind == "attachment" else 0,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )
            if kind == "attachment":
                actions.append("按附件格式设置为左空2字、首行缩进2字，正文下空1行")
                if text.endswith(("。", "；", "，")):
                    warnings.append("附件名称后建议不加标点符号。")
            else:
                actions.append("按正文层级设置为方正仿宋_GBK三号，首行缩进2字")
        elif kind == "salutation":
            _apply_paragraph_style(
                paragraph,
                BODY_FONT,
                BODY_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )
            actions.append("按称呼段设置为方正仿宋_GBK三号顶格")
        elif kind == "date":
            _apply_paragraph_style(
                paragraph,
                LEVEL2_FONT,
                BODY_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            actions.append("按年月设置为方正楷体_GBK三号居中")
            if any(ch in text for ch in "〇零壹贰叁肆伍陆柒捌玖拾"):
                warnings.append("成文日期疑似使用大写汉字，建议改为阿拉伯数字。")
        elif kind == "signature_date":
            _apply_paragraph_style(
                paragraph,
                BODY_FONT,
                BODY_SIZE,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )
            actions.append("按落款日期设置为方正仿宋_GBK三号右对齐")
        else:
            _apply_paragraph_style(
                paragraph,
                BODY_FONT,
                BODY_SIZE,
                bold=False,
                first_line_chars=2,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )
            actions.append("按正文设置为方正仿宋_GBK三号，首行缩进2字")

        _apply_western_font_normalization(paragraph, actions)
        _apply_keypoint_bold(paragraph, kind)
        _collect_sequence_warnings(kind, text, warnings)
        if len(text) > 120 and kind in {"title", "level1", "level2"}:
            warnings.append("该段较长但被识别为标题类，建议人工复核。")

        report = _build_paragraph_report(paragraph, idx + 1, kind, text, actions, warnings)
        summary.paragraph_reports.append(report)
        summary.warnings.extend(f"第{report.index}段：{warning}" for warning in warnings)
        summary.paragraphs_updated += 1


def _classify_paragraph(
    text: str,
    index: int,
    title_indices: set[int],
    author_indices: set[int],
    signature_indices: set[int] | None = None,
    signature_date_index: int | None = None,
) -> str:
    if DOC_NO_RE.match(text):
        return "doc_number"
    if signature_date_index is not None and index == signature_date_index:
        return "signature_date"
    if signature_indices and index in signature_indices:
        return "signature"
    if index in author_indices:
        return "author"
    if index in title_indices:
        return "title"
    if LEVEL1_RE.match(text):
        return "level1"
    if LEVEL2_RE.match(text):
        return "level2"
    if LEVEL3_RE.match(text):
        return "level3"
    if LEVEL4_RE.match(text):
        return "level4"
    if ATTACHMENT_RE.match(text):
        return "attachment"
    if _is_salutation_text(text):
        return "salutation"
    if DATE_RE.match(text):
        return "date"
    return "body"


def _detect_title_indices(paragraphs: list[Paragraph]) -> set[int]:
    non_empty = [(idx, p.text.strip()) for idx, p in enumerate(paragraphs) if p.text.strip()]
    if not non_empty:
        return set()

    doc_no_index = next((idx for idx, text in non_empty if DOC_NO_RE.match(text)), None)
    salutation_index = next((idx for idx, text in non_empty if _is_salutation_text(text)), None)
    first_level1_index = next((idx for idx, text in non_empty if LEVEL1_RE.match(text)), None)

    start = 0
    if doc_no_index is not None:
        start = next((pos for pos, (idx, _) in enumerate(non_empty) if idx == doc_no_index), 0) + 1

    end = None
    if salutation_index is not None:
        end = next((pos for pos, (idx, _) in enumerate(non_empty) if idx == salutation_index), None)
    elif first_level1_index is not None:
        end = next((pos for pos, (idx, _) in enumerate(non_empty) if idx == first_level1_index), None)

    candidates = non_empty[start:end]
    if not candidates and end is not None:
        candidates = non_empty[:end]

    title_indices = set()
    for idx, text in candidates[:2]:
        if (
            len(text) <= 40
            and not re.search(r"[。；，：:]", text)
            and not re.fullmatch(rf"[{ZH_CHAR_RE}·]{{2,8}}", text)
            and not DATE_RE.match(text)
            and not LEVEL1_RE.match(text)
        ):
            title_indices.add(idx)

    if not title_indices and non_empty:
        first_idx, first_text = non_empty[0]
        if len(first_text) <= 40 and first_level1_index not in {None, first_idx}:
            title_indices.add(first_idx)

    return title_indices


def _detect_author_indices(paragraphs: list[Paragraph], title_indices: set[int]) -> set[int]:
    non_empty = [(idx, p.text.strip()) for idx, p in enumerate(paragraphs) if p.text.strip()]
    if not non_empty or not title_indices:
        return set()

    first_title_index = min(title_indices)
    first_level1_index = next((idx for idx, text in non_empty if LEVEL1_RE.match(text)), None)

    author_indices: set[int] = set()
    for idx, text in non_empty:
        if idx <= first_title_index:
            continue
        if first_level1_index is not None and idx >= first_level1_index:
            break
        if DATE_RE.match(text):
            continue
        if (
            len(text) <= 8
            and re.fullmatch(rf"[{ZH_CHAR_RE}·]{{2,8}}", text)
            and not text.endswith(("：", ":"))
        ):
            author_indices.add(idx)
            break
    return author_indices


def _is_salutation_text(text: str) -> bool:
    excluded_keywords = ("要求", "流程", "说明", "事项", "如下", "原则", "标准", "条件")
    return bool(
        text.endswith(("：", ":"))
        and len(text) <= 24
        and not re.search(r"[。；，！？]", text)
        and not LEVEL1_RE.match(text)
        and not LEVEL2_RE.match(text)
        and not LEVEL3_RE.match(text)
        and not LEVEL4_RE.match(text)
        and not any(keyword in text for keyword in excluded_keywords)
    )


def _detect_signature_indices(paragraphs: list[Paragraph]) -> tuple[set[int], int | None]:
    non_empty = [(idx, p.text.strip()) for idx, p in enumerate(paragraphs) if p.text.strip()]
    if not non_empty:
        return set(), None

    date_idx = None
    for idx, text in reversed(non_empty):
        if SIGNATURE_DATE_RE.match(text):
            date_idx = idx
            break
    if date_idx is None:
        return set(), None

    signature_indices: set[int] = set()
    taken = 0
    for idx, text in reversed(non_empty):
        if idx >= date_idx:
            continue
        if taken >= 3:
            break
        if (
            len(text) <= 24
            and not re.search(r"[。；，：:]", text)
            and not LEVEL1_RE.match(text)
            and not LEVEL2_RE.match(text)
            and not LEVEL3_RE.match(text)
            and not LEVEL4_RE.match(text)
            and not ATTACHMENT_RE.match(text)
            and not DOC_NO_RE.match(text)
            and not text.endswith(("：", ":"))
        ):
            signature_indices.add(idx)
            taken += 1
            continue
        break

    return signature_indices, date_idx


def _apply_paragraph_style(
    paragraph: Paragraph,
    east_asia_font: str,
    font_size,
    *,
    bold: bool,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    first_line_chars: int | None = None,
    left_chars: int | None = None,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    _normalize_paragraph_spacing(paragraph)
    if alignment is not None:
        paragraph.alignment = alignment
    _clear_paragraph_border(paragraph)

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = LINE_SPACING
    paragraph_format.space_before = Pt(space_before_pt)
    paragraph_format.space_after = Pt(space_after_pt)
    paragraph_format.left_indent = Pt(BODY_SIZE.pt * left_chars) if left_chars is not None else Pt(0)
    paragraph_format.first_line_indent = Pt(BODY_SIZE.pt * first_line_chars) if first_line_chars is not None else Pt(0)
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.keep_together = keep_together

    if not paragraph.runs:
        run = paragraph.add_run("")
        _set_run_fonts(run, east_asia_font, font_size, bold)
        return

    for run in paragraph.runs:
        _set_run_fonts(run, east_asia_font, font_size, bold)


def _set_run_fonts(run: Run, east_asia_font: str, font_size, bold: bool) -> None:
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = WESTERN_FONT
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.highlight_color = None
    run.font.underline = False

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:ascii"), WESTERN_FONT)
    r_fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    r_fonts.set(qn("w:cs"), WESTERN_FONT)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    color = r_pr.color
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), "000000")


def _normalize_title_text(paragraph: Paragraph) -> None:
    text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    text = text.strip()
    if not text:
        return

    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    lines = _split_title_lines(text)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        if idx < len(lines) - 1:
            run.add_break()


def _split_title_lines(text: str) -> list[str]:
    if len(text) <= 22:
        return [text]

    if "（" in text and text.endswith("）"):
        main, suffix = text.split("（", 1)
        main = main.strip()
        suffix = "（" + suffix.strip()
        if len(main) >= 10:
            first, second = _pick_title_break(main)
            return [first, second + suffix]

    first, second = _pick_title_break(text)
    return [first, second]


def _pick_title_break(text: str) -> tuple[str, str]:
    midpoint = len(text) // 2
    preferred_break_before = [
        "征求意见汇总",
        "工作学习心得",
        "学习心得",
        "情况汇报",
        "工作总结",
        "实施方案",
        "征求意见",
        "意见汇总",
        "工作要点",
        "总结",
        "汇总",
        "方案",
        "报告",
        "通知",
    ]
    candidates = []
    for phrase in preferred_break_before:
        pos = text.find(phrase)
        if pos != -1 and 8 <= pos <= len(text) - 8:
            candidates.append(pos)

    if candidates:
        break_pos = min(candidates, key=lambda pos: abs(pos - midpoint))
        return text[:break_pos], text[break_pos:]

    semantic_break_after = [
        "组织生活会征求意见汇总",
        "组织生活会",
        "工作学习心得",
        "征求意见汇总",
        "征求意见",
        "意见汇总",
        "年度",
        "会议",
        "情况",
        "报告",
        "通知",
    ]
    trailing = []
    for phrase in semantic_break_after:
        pos = text.find(phrase)
        if pos != -1:
            break_pos = pos + len(phrase)
            if 8 <= break_pos <= len(text) - 8:
                trailing.append(break_pos)
    if trailing:
        break_pos = min(trailing, key=lambda pos: abs(pos - midpoint))
        return text[:break_pos], text[break_pos:]

    candidates = []
    for i in range(8, len(text) - 8):
        if text[i - 1] not in "的一了和与及并" and text[i] not in "的一了和与及并":
            candidates.append(i)
    break_pos = min(candidates or [midpoint], key=lambda pos: abs(pos - midpoint))
    return text[:break_pos], text[break_pos:]


def _clear_paragraph_border(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_pr.remove(p_style)
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right", "between"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "nil")
        p_bdr.append(border)
    p_pr.append(p_bdr)


def _normalize_paragraph_spacing(paragraph: Paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = LINE_SPACING
    if paragraph.alignment is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _apply_western_font_normalization(paragraph: Paragraph, actions: list[str]) -> None:
    if LATIN_RE.search(paragraph.text):
        actions.append("数字和西文字母统一设置为 Times New Roman")


def _apply_keypoint_bold(paragraph: Paragraph, kind: str) -> None:
    if kind not in {"body", "salutation", "level3", "level4", "attachment"}:
        return
    text = paragraph.text
    if not KEYPOINT_RE.search(text):
        return

    segments: list[tuple[str, bool]] = []
    cursor = 0
    for match in KEYPOINT_RE.finditer(text):
        if match.start() > cursor:
            segments.append((text[cursor:match.start()], False))
        segments.append((match.group(0), True))
        cursor = match.end()
    if cursor < len(text):
        segments.append((text[cursor:], False))

    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    for content, is_bold in segments:
        run = paragraph.add_run(content)
        _set_run_fonts(run, BODY_FONT, BODY_SIZE, is_bold)


def _collect_sequence_warnings(kind: str, text: str, warnings: list[str]) -> None:
    if kind == "level1" and not LEVEL1_RE.match(text):
        warnings.append("一级标题序号格式可能不符合“一、”规范。")
    if kind == "level2" and not LEVEL2_RE.match(text):
        warnings.append("二级标题序号格式可能不符合“（一）”规范。")
    if kind == "level3" and not LEVEL3_RE.match(text):
        warnings.append("三级标题序号格式可能不符合“1.”规范。")
    if kind == "level4" and not LEVEL4_RE.match(text):
        warnings.append("四级标题序号格式可能不符合“（1）”规范。")
    if kind == "body" and re.match(r"^[（(]?[一二三四五六七八九十\d]+[）).、]", text):
        warnings.append("该段看起来像有层次序号，但未被识别到对应级别，建议人工复核。")


def _build_paragraph_report(
    paragraph: Paragraph,
    index: int,
    kind: str,
    text: str,
    actions: list[str],
    warnings: list[str],
) -> ParagraphReport:
    run = next((item for item in paragraph.runs if item.text), paragraph.runs[0] if paragraph.runs else None)
    east_font = ""
    ascii_font = ""
    font_size_pt = 0.0
    if run is not None:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is not None:
            east_font = r_fonts.get(qn("w:eastAsia")) or ""
            ascii_font = r_fonts.get(qn("w:ascii")) or ""
        if run.font.size is not None:
            font_size_pt = float(run.font.size.pt)

    pf = paragraph.paragraph_format
    line_spacing_pt = float(pf.line_spacing.pt) if hasattr(pf.line_spacing, "pt") else float(LINE_SPACING.pt)
    return ParagraphReport(
        index=index,
        kind=kind,
        text=text,
        font_east_asia=east_font,
        font_ascii=ascii_font,
        font_size_pt=font_size_pt,
        alignment=_alignment_name(paragraph.alignment),
        first_line_indent_pt=float(pf.first_line_indent.pt) if pf.first_line_indent is not None else 0.0,
        left_indent_pt=float(pf.left_indent.pt) if pf.left_indent is not None else 0.0,
        line_spacing_pt=line_spacing_pt,
        actions=actions,
        warnings=warnings,
    )


def _alignment_name(alignment: WD_ALIGN_PARAGRAPH | None) -> str:
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(alignment, "unknown")


def _format_tables(doc: DocumentObject, summary: FormatSummary) -> None:
    for table in doc.tables:
        _format_table(table)
        summary.tables_updated += 1
    if doc.tables:
        summary.warnings.append("表格内容已统一基础字体，但未对表格边框、列宽、跨页策略做专项优化。")


def _format_table(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _format_cell(cell)


def _format_cell(cell: _Cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        kind = "level1" if LEVEL1_RE.match(text) else "body"
        if kind == "level1":
            _apply_paragraph_style(paragraph, LEVEL1_FONT, BODY_SIZE, bold=False, first_line_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _apply_paragraph_style(paragraph, BODY_FONT, BODY_SIZE, bold=False, first_line_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for table in cell.tables:
        _format_table(table)


def _set_page_footer(doc: DocumentObject) -> bool:
    doc.settings.odd_and_even_pages_header_footer = True
    updated = False
    for section in doc.sections:
        _write_page_footer(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
        _write_page_footer(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)
        updated = True
    return updated


def _write_page_footer(footer, alignment: WD_ALIGN_PARAGRAPH) -> None:
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = alignment
    paragraph.clear()
    run_left = paragraph.add_run("—")
    _set_page_fonts(run_left)
    _append_page_field(paragraph)
    run_right = paragraph.add_run("—")
    _set_page_fonts(run_right)


def _set_page_fonts(run: Run) -> None:
    run.font.size = PAGE_SIZE
    run.font.name = PAGE_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), WESTERN_FONT)
    r_fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    r_fonts.set(qn("w:cs"), WESTERN_FONT)
    r_fonts.set(qn("w:eastAsia"), PAGE_FONT)


def _append_page_field(paragraph: Paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    _set_page_fonts(run)
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
